from __future__ import annotations

from io import BytesIO
from datetime import datetime
from typing import Dict, Any
import pandas as pd
from pandas.api.types import is_scalar

from docx import Document


def _cell_to_text(val: object) -> str:
    """Convert arbitrary pandas cell values to safe text.

    - Handles duplicate column names (row[c] -> Series)
    - Handles NA values safely
    - Produces deterministic strings for list/dict-like objects
    """
    if isinstance(val, pd.Series):
        cleaned = []
        for v in val.tolist():
            if is_scalar(v) and pd.isna(v):
                continue
            cleaned.append(v)
        return "" if not cleaned else ", ".join(str(v) for v in cleaned)

    if is_scalar(val):
        return "" if pd.isna(val) else str(val)

    try:
        na_mask = pd.isna(val)
        if hasattr(na_mask, "all") and bool(na_mask.all()):
            return ""
    except Exception:
        pass

    return str(val)


def _safe_title(val: object, fallback: str = "Report") -> str:
    txt = _cell_to_text(val).strip()
    return txt if txt else fallback


def build_summary(df: pd.DataFrame) -> Dict[str, Any]:
    """Fast, safe executive summary.

    Returns ONLY Python scalars/strings (no pandas objects) to avoid downstream UI/export issues.
    """
    if df is None or df.empty:
        return {"Total rows": 0, "Total columns": 0}

    total_rows = int(len(df))
    total_cols = int(len(df.columns))

    # Avoid expensive full astype(str) across huge tables.
    # Sample up to 80 columns for global cell stats.
    sample_cols = list(df.columns[: min(80, total_cols)])

    nonempty_cells = 0
    empty_cells = 0
    for c in sample_cols:
        s = df[c]
        # treat NA + blank strings as empty
        if s.dtype == object:
            ss = s.astype(str).replace({"nan": "", "None": ""}).str.strip()
            ne = int((ss != "").sum())
        else:
            ne = int(s.notna().sum())
        nonempty_cells += ne
        empty_cells += total_rows - ne

    # Normalize to estimated totals if we sampled columns
    if total_cols > len(sample_cols):
        scale = total_cols / max(1, len(sample_cols))
        nonempty_cells = int(round(nonempty_cells * scale))
        empty_cells = int(round(empty_cells * scale))

    summary: Dict[str, Any] = {
        "Total rows": total_rows,
        "Total columns": total_cols,
        "Non-empty cells (est.)": int(nonempty_cells),
        "Missing cells (est.)": int(empty_cells),
    }

    # A few useful column-specific KPIs if present (safe + cheap)
    for col in ["Confirm gender", "_gender_norm", "Province", "District", "Village name:", "Age group"]:
        if col in df.columns:
            s = df[col].astype(str).replace({"nan": "", "None": ""}).str.strip()
            summary[f"Non-empty: {col}"] = int((s != "").sum())

    if "_uuid" in df.columns:
        s = df["_uuid"].astype(str).replace({"nan": "", "None": ""}).str.strip()
        summary["UUID filled"] = int((s != "").sum())
        try:
            summary["UUID duplicates"] = int(df.duplicated(subset=["_uuid"]).sum())
        except Exception:
            summary["UUID duplicates"] = 0

    return summary


def to_csv_bytes(df: pd.DataFrame) -> bytes:
    return df.to_csv(index=False).encode("utf-8-sig")


def to_excel_bytes(df: pd.DataFrame, summary: Dict[str, Any]) -> bytes:
    bio = BytesIO()
    with pd.ExcelWriter(bio, engine="openpyxl") as writer:
        df.to_excel(writer, sheet_name="data", index=False)

        s = pd.DataFrame({"Metric": list(summary.keys()), "Value": list(summary.values())})
        s.to_excel(writer, sheet_name="summary", index=False)

    return bio.getvalue()


def to_word_bytes(tool_name: object, df: pd.DataFrame, summary: Dict[str, Any]) -> bytes:
    doc = Document()
    tool_name_txt = _safe_title(tool_name, fallback="UN Women")
    doc.add_heading(f"Auto Report - {tool_name_txt}", level=1)
    doc.add_paragraph(f"Generated at: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")  # UTC

    doc.add_heading("Summary", level=2)
    for k, v in summary.items():
        doc.add_paragraph(f"{k}: {v}")

    doc.add_heading("Data Preview (first 15 rows)", level=2)
    preview = df.head(15)

    if preview.empty:
        doc.add_paragraph("No rows available.")
    else:
        table = doc.add_table(rows=1, cols=len(preview.columns))
        hdr = table.rows[0].cells
        for j, c in enumerate(preview.columns):
            hdr[j].text = str(c)

        for _, row in preview.iterrows():
            cells = table.add_row().cells
            for j, c in enumerate(preview.columns):
                cells[j].text = _cell_to_text(row[c])

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
